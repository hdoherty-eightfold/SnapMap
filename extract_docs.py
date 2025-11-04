from docx import Document
import os

docs_folder = r'C:\Code\SnapMap\docs'
output_folder = r'C:\Code\SnapMap\docs_extracted'

os.makedirs(output_folder, exist_ok=True)

doc_files = [
    '00_START_HERE_Reading_Order.docx',
    '01_Quick_Summary_5min.docx',
    '02_Master_Specification_MAIN.docx',
    '03_README_Quick_Start.docx',
    '04_Week_Plan_Daily_Tasks.docx',
    '05_Starter_Code_and_Cursor_Prompts.docx',
    '09_Modular_Development_Plan.docx',
    '10_API_Contracts_Template.docx',
    'REF-1_Full_Solution_Specifications.docx',
    'REF-2_Implementation_Guide.docx'
]

for doc_file in doc_files:
    try:
        doc_path = os.path.join(docs_folder, doc_file)
        if not os.path.exists(doc_path):
            print(f"Skipping {doc_file} - not found")
            continue

        doc = Document(doc_path)

        output_file = os.path.join(output_folder, doc_file.replace('.docx', '.txt'))

        with open(output_file, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                f.write(para.text + '\n')

            # Also extract tables
            for table in doc.tables:
                f.write('\n--- TABLE ---\n')
                for row in table.rows:
                    row_text = '\t'.join([cell.text for cell in row.cells])
                    f.write(row_text + '\n')
                f.write('--- END TABLE ---\n\n')

        print(f"Extracted: {doc_file}")
    except Exception as e:
        print(f"Error extracting {doc_file}: {e}")

print("\nExtraction complete!")
